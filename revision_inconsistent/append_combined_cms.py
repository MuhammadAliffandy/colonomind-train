"""
Gabungkan semua CM image per tabel, lalu ganti section gambar di DOCX.
Script ini MENGGANTI (bukan menambahkan) halaman Combined CM yang lama.
"""
import os
from PIL import Image, ImageDraw, ImageFont
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE   = "/Users/aliffandy/Documents/PukulEnam/Colonomind Training Resource/revision_inconsistent"
CM_DIR = os.path.join(BASE, "cm_images_v2")
DOC_IN = os.path.join(BASE, "ColonoMind_Table1_CM_Final.docx")

SECTION_MARKER = "==COMBINED_CM_SECTION=="   # penanda tersembunyi di paragraf pertama section

def combine_images(files, title, out_filename, n_cols=3):
    images = []
    for f in files:
        path = os.path.join(CM_DIR, f)
        if os.path.exists(path):
            images.append(Image.open(path))
        else:
            print(f"  MISSING: {f}")

    if not images:
        return None

    w, h   = images[0].size
    pad    = 40
    title_h = 90
    n_rows = -(-len(images) // n_cols)          # ceiling div

    total_w = n_cols * w + (n_cols + 1) * pad
    total_h = n_rows * h + (n_rows + 1) * pad + title_h

    canvas = Image.new('RGB', (total_w, total_h), 'white')
    draw   = ImageDraw.Draw(canvas)

    try:
        font_title = ImageFont.truetype("Arial Bold.ttf", 54)
    except:
        try:
            font_title = ImageFont.truetype("Arial", 54)
        except:
            font_title = ImageFont.load_default()

    bbox = draw.textbbox((0, 0), title, font=font_title)
    tw   = bbox[2] - bbox[0]
    draw.text(((total_w - tw) / 2, pad), title, fill='black', font=font_title)

    for idx, img in enumerate(images):
        row = idx // n_cols
        col = idx % n_cols
        # Center last incomplete row
        imgs_this_row = min(n_cols, len(images) - row * n_cols)
        row_width     = imgs_this_row * w + (imgs_this_row - 1) * pad
        x_start       = (total_w - row_width) // 2
        x = x_start + col * (w + pad)
        y = title_h + pad + row * (h + pad)
        canvas.paste(img, (x, y))

    out_path = os.path.join(BASE, out_filename)
    canvas.save(out_path, dpi=(300, 300))
    print(f"  Saved → {out_filename}")
    return out_path


def remove_old_cm_section(doc):
    """Hapus semua elemen setelah penanda SECTION_MARKER."""
    body = doc.element.body
    paragraphs = doc.paragraphs

    marker_idx = None
    for i, para in enumerate(paragraphs):
        if SECTION_MARKER in para.text:
            marker_idx = i
            break

    if marker_idx is None:
        return False    # tidak ada section lama, tidak perlu hapus

    # Kumpulkan semua elemen XML dari marker ke akhir body
    all_body_children = list(body)
    para_xml_list = [para._element for para in paragraphs]

    marker_xml = para_xml_list[marker_idx]
    found = False
    to_remove = []
    for child in all_body_children:
        if child == marker_xml:
            found = True
        if found:
            to_remove.append(child)

    for el in to_remove:
        body.remove(el)

    print(f"  Removed old CM section ({len(to_remove)} elements)")
    return True


def add_hidden_marker(doc):
    """Tambahkan paragraf penanda tersembunyi (warna putih) untuk deteksi saat re-run."""
    para = doc.add_paragraph(SECTION_MARKER)
    run  = para.runs[0]
    run.font.color.rgb = docx.shared.RGBColor(0xFF, 0xFF, 0xFF)  # putih = tidak terlihat
    run.font.size = Pt(1)


def run():
    all_files = sorted(f for f in os.listdir(CM_DIR) if f.endswith('.png'))

    tables = {
        "T1": ([f for f in all_files if f.startswith("T1_")],
               "Table 1 — Intra-domain confusion matrices",
               "Combined_T1.png", 3),
        "T2": ([f for f in all_files if f.startswith("T2_")],
               "Table 2 — Cross-domain confusion matrices",
               "Combined_T2.png", 3),
        "T3": ([f for f in all_files if f.startswith("T3_")],
               "Table 3 — Multi-domain (LIMUC) confusion matrices",
               "Combined_T3.png", 3),
        "T4": ([f for f in all_files if f.startswith("T4_")],
               "Table 4 — Multi-domain (TMC-UCM) confusion matrices",
               "Combined_T4.png", 3),
    }

    print("═" * 60)
    print("STEP 1: Menggabungkan gambar CM per tabel...")
    combined_paths = []
    for tkey, (files, title, fname, n_cols) in tables.items():
        print(f"\n[{tkey}] {title}")
        p = combine_images(files, title, fname, n_cols)
        if p:
            combined_paths.append((title, p))

    print("\n═" * 60)
    print("STEP 2: Update DOCX — menghapus halaman CM lama dan memasukkan yang baru...")

    doc = docx.Document(DOC_IN)
    removed = remove_old_cm_section(doc)
    if not removed:
        print("  (Tidak ada section lama — gambar langsung ditambahkan di halaman baru)")

    # Tambah page break + marker tersembunyi
    doc.add_page_break()
    add_hidden_marker(doc)

    # Judul section
    heading = doc.add_heading("Confusion Matrices — All Tables", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Masukkan tiap gambar di halaman terpisah
    for i, (title, img_path) in enumerate(combined_paths):
        sub = doc.add_heading(title, level=2)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        run.add_picture(img_path, width=Inches(6.8))
        if i < len(combined_paths) - 1:
            doc.add_page_break()

    doc.save(DOC_IN)
    print(f"\n✅ Selesai! DOCX diperbarui: {os.path.basename(DOC_IN)}")
    print(f"   Gambar gabungan: {len(combined_paths)} tabel, masing-masing di halaman sendiri.")

if __name__ == "__main__":
    run()
